from docx import Document
import os
import shutil
import tempfile
import traceback
import zipfile

def docx_formating(source_file_path, output_folder):
    try:
        # Create a temporary directory within the specified output folder to extract images
        with tempfile.TemporaryDirectory(dir=output_folder) as temp_extract_dir:
            # Extract the content of the source DOCX file
            with zipfile.ZipFile(source_file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_extract_dir)

            # Read the extracted content
            extracted_files = os.listdir(temp_extract_dir)

            new_doc = Document()
            image_src_map = {}  # Map of image source paths to new image paths

            for file_name in extracted_files:
                file_path = os.path.join(temp_extract_dir, file_name)

                # Handle images
                if file_name.startswith("word/media/"):
                    new_image_path = os.path.join(output_folder, file_name[11:])  # Removing 'word/media/'
                    shutil.copy(file_path, new_image_path)
                    image_src_map[file_name] = new_image_path
                else:
                    # Handle other content
                    print ( "file path ", file_path)
                    with open(file_path, 'r', encoding='utf-8') as file:
                        content = file.read()
                        new_doc.add_paragraph(content)

            # Replace image source paths in the content
            for p in new_doc.paragraphs:
                for r in p.runs:
                    for old_src, new_src in image_src_map.items():
                        if old_src in r._element.xml:
                            r._element.xml = r._element.xml.replace(old_src, new_src)

            # Save the new formatted DOCX file
            new_docx_path = os.path.join(output_folder, os.path.basename(source_file_path))
            new_doc.save(new_docx_path)

            return new_docx_path

    except Exception as e:
        print(f"Error in docx_formating: {e}")
        traceback.print_exc()
        return None

# Example usage
source_file_path = "E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\docx_formatinng\ok.docx"
output_folder = "E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\docx_formatinng\\test"
new_docx_path = docx_formating(source_file_path, output_folder)
if new_docx_path:
    print(f"Formatting completed. New DOCX saved at: {new_docx_path}")
else:
    print("Formatting failed.")