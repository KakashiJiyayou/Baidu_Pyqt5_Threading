import os
import sys

from docx import Document
from PIL import Image
import re

def find_image_path(image_path):
    if os.path.exists(image_path):
        return image_path
    return None


def process_docx(docx_path, output_folder , output_docx_path):
    document = Document(docx_path)
    modified_lines = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        # img_matches = re.findall(r'IMAG_SRC_\d+@@@.*', line)
        # print ( "img_matches" , img_matches  )

        if line.startswith("IMAG_SRC_"):
            # sys.exit()
            try:
                print ( line.split("@@@") )
                img_placeholder, img_path = line.split("@@@")
                real_image_path = find_image_path(img_path)
            except:
                real_image_path = None
                print ("error")

            if real_image_path:
                img = Image.open(real_image_path)
                img = img.convert('RGB')

                # Set the desired aspect ratio (height-to-width ratio)
                desired_aspect_ratio = 0.75  # You can adjust this value

                # Calculate the new width based on the standard height and desired aspect ratio
                standard_height = 300  # You can adjust this value
                new_width = int(standard_height / desired_aspect_ratio)

                # Resize the image to the new dimensions (standard_height x new_width)
                img = img.resize((new_width, standard_height), Image.ANTIALIAS)

                temp_image_path = os.path.join(output_folder, 'temp_image.jpg')
                img.save(temp_image_path, format='JPEG')
                paragraph.clear()  # Clear the current paragraph's content
                run = paragraph.add_run()  # Add a run to the paragraph
                run.add_picture(temp_image_path)  # Add the resized image
                os.remove(temp_image_path)  # Remove the temporary image file
            else:
                line += "_NOT_FOUND"

        modified_lines.append(line)
        document.save ( output_docx_path )


# Output folder for temporary images
output_folder = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\final_test"

# Input DOCX path
input_docx_path = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\libre\docx_from_html\tempx1.docx"

base_name = os.path.basename( input_docx_path )
output_docx_path =  os.path.join( output_folder, base_name )

process_docx(input_docx_path, output_folder , output_docx_path)


print("Processing completed successfully.")