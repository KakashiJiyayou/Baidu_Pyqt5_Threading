import os
from docx import Document
from PIL import Image


def find_image_path(image_path):
    if os.path.exists(image_path):
        return image_path
    return None


def process_docx(docx_path, output_folder , output_docx_path):
    document = Document(docx_path)
    modified_lines = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line.startswith("IMAG_SRC_"):
            img_placeholder, img_path = line.split("@@@")
            real_image_path = find_image_path(img_path)

            if real_image_path:
                img = Image.open(real_image_path)
                img = img.convert('RGB')

                # Set the desired aspect ratio (height-to-width ratio)
                desired_aspect_ratio = 0.75  # You can adjust this value

                # Calculate the new width based on the standard height and desired aspect ratio
                standard_height = 300  # You can adjust this value
                new_width = int(standard_height / desired_aspect_ratio)

                # Resize the image to the new dimensions (standard_height x new_width)
                img = img.resize((new_width, standard_height), Image.ANTIALIAS)

                temp_image_path = os.path.join(output_folder, 'temp_image.jpg')
                img.save(temp_image_path, format='JPEG')
                paragraph.clear()  # Clear the current paragraph's content
                run = paragraph.add_run()  # Add a run to the paragraph
                run.add_picture(temp_image_path)  # Add the resized image
                os.remove(temp_image_path)  # Remove the temporary image file
            else:
                line += "_NOT_FOUND"

        modified_lines.append(line)
        document.save ( output_docx_path )







# Input DOCX path
input_docx_path = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\add_image_to_Doc\CCTV_m1.docx"

# Output DOCX path for saving modified content
output_docx_path = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\add_image_to_Doc\Modified_CCTV_m1.docx"

# Output folder for temporary images
output_folder = "E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\\add_image_to_Doc"

# Process the DOCX file
modified_lines = process_docx(input_docx_path, output_folder , output_docx_path)


print("Processing completed successfully.")