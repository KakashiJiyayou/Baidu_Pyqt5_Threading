import os
from docx import Document


def find_image_path(image_path):
    if os.path.exists(image_path):
        return image_path
    return None


def process_docx(docx_path):
    document = Document(docx_path)
    modified_lines = []

    for paragraph in document.paragraphs:
        line = paragraph.text.strip()
        if line.startswith("IMAG_SRC_"):
            img_placeholder, img_path = line.split("@@@")
            real_image_path = find_image_path(img_path)

            if real_image_path:
                # Add the image without specifying width
                document.add_picture(real_image_path)
            else:
                line += "_NOT_FOUND"

        modified_lines.append(line)

    return modified_lines


def save_modified_docx(lines, output_path):
    document = Document()
    for line in lines:
        document.add_paragraph(line)
    document.save(output_path)


# Input DOCX path
input_docx_path = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\add_image_to_Doc\CCTV_m1.docx"

# Output DOCX path for saving modified content
output_docx_path = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_to_html\add_image_to_Doc\Modified_CCTV_m1.docx"

# Process the DOCX file
modified_lines = process_docx(input_docx_path)

# Save the modified content to a new DOCX file
save_modified_docx(modified_lines, output_docx_path)

print("Processing completed successfully.")
