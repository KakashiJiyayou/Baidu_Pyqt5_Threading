from docx import Document
from PIL import Image
from io import BytesIO


def extract_and_merge_images(doc_paths, output_path):
    merged_doc = Document()

    for doc_path in doc_paths:
        doc = Document(doc_path)

        for element in doc.element.body:
            if element.tag.endswith('r'):
                for child in element:
                    if child.tag.endswith('drawing'):
                        for grandchild in child:
                            if grandchild.tag.endswith('blip'):
                                img_part = grandchild.get_or_add('a:blip')
                                img_id = img_part.get('r:embed')
                                img = doc.part.related_parts[img_id].blob
                                img = Image.open(BytesIO(img))

                                img_bytes = BytesIO()
                                img.save(img_bytes, format='PNG')
                                merged_doc.add_picture(img_bytes, width=None, height=None)
                                merged_doc.add_paragraph()  # Add an empty paragraph after each image

    merged_doc.save(output_path)
    print(f"Merged document with images saved to {output_path}")



# List of input docx files
input_docx_paths = [
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_0.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_1.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_2.docx',
    # Add more paths as needed
]

# Output path for the merged document
output_docx_path = 'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\gen\\new2.docx'


# Call the function to extract images and merge the documents
extract_and_merge_images(input_docx_paths, output_docx_path)