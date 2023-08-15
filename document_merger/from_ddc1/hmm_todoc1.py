from docx import Document

import os

def extract_images_from_docx(docx_path, output_folder):
    doc = Document(docx_path)
    image_counter = 1

    for shape in doc.inline_shapes:
        if shape.type == 3:  # Check if it's an image
            image_data = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_path = os.path.join(output_folder, f'image{image_counter}.png')
            with open(image_path, 'wb') as img_file:
                img_file.write(doc.part.related_parts[image_data].blob)
            image_counter += 1

def merge_docx_files(doc_paths, output_path):
    merged_doc = Document()

    for doc_path in doc_paths:
        doc = Document(doc_path)

        for paragraph in doc.paragraphs:
            merged_doc.add_paragraph(paragraph.text)

    merged_doc.save(output_path)
    print(f"Merged documents saved to {output_path}")

# List of input docx files
files_list = [
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_0.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_1.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_2.docx',
    # Add more paths as needed
]

# Output folder for images and merged document
output_folder = 'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\gen'

# Output filename for the merged document
output_filename = 'merged.docx'

# Call the function to extract images from the docx files
for doc_path in files_list:
    extract_images_from_docx(doc_path, output_folder)

# Call the function to merge the documents
output_path = os.path.join(output_folder, output_filename)
merge_docx_files(files_list, output_path)
