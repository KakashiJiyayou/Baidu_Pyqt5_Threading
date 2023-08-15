from docx import Document


def merge_docx_files(doc_paths, output_path):
    merged_doc = Document()

    for doc_path in doc_paths:
        doc = Document(doc_path)

        for element in doc.element.body:
            if element.tag.endswith('p'):
                new_paragraph = Document()
                new_paragraph.element.body.append(element)
                merged_doc.add_paragraph(new_paragraph.paragraphs[0].text, style=merged_doc.styles['BodyText'])

    merged_doc.save(output_path)
    print(f"Merged documents saved to {output_path}")


# List of input docx files
input_docx_paths = [
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_0.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_1.docx',
    'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\\temp\\temp_doc_2.docx',
    # Add more paths as needed
]

# Output path for the merged document
output_docx_path = 'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\\from_ddc1\gen\\new2.docx'

# Call the function to merge the documents
merge_docx_files(input_docx_paths, output_docx_path)