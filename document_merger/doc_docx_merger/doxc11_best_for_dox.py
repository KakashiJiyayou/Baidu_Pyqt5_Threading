from docxcompose.composer import Composer
from docx import Document

import os
def merge_documents(doc_paths, output_path):

    # master_doc1 = Document()
    # master_doc1.save(doc_paths[0])
    #
    #
    # master_doc = Document(doc_paths[0])
    #
    # composer = Composer(master_doc)
    #
    # for doc_path in doc_paths[1:]:
    #
    #     doc1 = Document()
    #     doc1.save(doc_path)
    #     # doc = Document(doc_path)
    #     doc = Document(doc_path)
    #     composer.append(doc)
    #
    # composer.save(output_path)

    master_doc = Document()
    master_doc.add_paragraph()

    # Create the composer with the master document
    composer = Composer(master_doc)

    for doc_path in doc_paths:
        try:
            doc = Document(doc_path)
            composer.append(doc)
        except:
            pass

    composer.save(output_path)



# List of paths to the documents you want to merge
input_files = [
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\1.营业执照.doc',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\2.市政资质.doc',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\CCTV检测技术文件10页.doc',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\CCTV检测文件5页.doc',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\下水管道养护工 张明华.doc',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\安全员于娜娜0725.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\附件2：2023 宁波市国际学生汉语口语大赛报名表.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\管道检测标应急预案.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\组织安全方案.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\安全员谢延杰0725.docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\组织安全方案.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\质量员程礼海(1).docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\2.市政资质.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\质量员程礼海0725.docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\资料员谢延莉(1).docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\2.市政资质.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\资料员谢延莉0725.docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\2.市政资质.doc',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\During 4th and 5th Competition.jpg',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\InternationalCompetition in Wuxi.jpg',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\Published Paper.jpg',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\screen v6.png',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\BildschirmÂ_foto 2022-11-16 um 16.12.00.png',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\附件2：2023 宁波市国际学生汉语口语大赛报名表.docx',
        # # Add more paths as needed
    ]

# Output file name for the merged document
output_file = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp\tue11.docx"

merge_documents(input_files, output_file)