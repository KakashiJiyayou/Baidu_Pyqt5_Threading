from docx import Document

def merge_docx_files(doc_paths, output_path):
    merged_doc = Document()

    for doc_path in doc_paths:
        doc = Document(doc_path)

        for paragraph in doc.paragraphs:
            merged_doc.add_paragraph(paragraph.text)

        for shape in doc.inline_shapes:
            if shape.type == 3:  # Check if it's an image
                image_data = shape.image.blob
                merged_doc.add_picture(image_data)

    merged_doc.save(output_path)
    print(f"Merged documents saved to {output_path}")

# List of input processed DOCX files
processed_docx_paths = [
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp1.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp2.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp3.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp4.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx1.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx2.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx3.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx4.docx',
    # Add more paths as needed
]

# Output path for the merged document
merged_output_path = r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp\new1.docx'

# Call the function to merge the processed documents
merge_docx_files(processed_docx_paths, merged_output_path)
