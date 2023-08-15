from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os


def copy_element_to_document(source_element, destination_doc):
    new_paragraph = destination_doc.add_paragraph()

    for run in source_element:
        new_run = new_paragraph.add_run(run.text)
        # new_run.bold = run.bold


def copy_images(source_doc, destination_doc):
    for rel in source_doc.part.rels:
        if "image" in source_doc.part.rels[rel].target_ref:
            image_part = source_doc.part.related_parts[rel]

            # Copy the image part from the source document to the destination document
            new_image_part = destination_doc.part.related_parts[rel]._blob

            # Create a new relationship for the copied image part in the destination document
            new_rel = destination_doc.part.relate_to(
                new_image_part, source_doc.part.rels[rel].target_ref
            )

            # Update the relationship ID in the destination document's body
            for paragraph in destination_doc.paragraphs:
                for run in paragraph.runs:
                    for elem in run._element:
                        if elem.tag.endswith('drawing'):
                            for inline in elem:
                                for prop in inline:
                                    if prop.tag.endswith('blip'):
                                        prop.attrib[
                                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'] = new_rel


# Rest of your code


def merge_docx(input_files, output_dir, output_filename):
    merged_doc = Document()

    for input_file in input_files:
        docx_content = Document(input_file)
        for element in docx_content.element.body:
            copy_element_to_document(element, merged_doc)

    copy_images(docx_content, merged_doc)

    # Save the merged document
    merged_path = os.path.join(output_dir, output_filename)
    merged_doc.save(merged_path)

if __name__ == "__main__":
    input_files = [
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp4.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx4.docx',
        # Add more paths as needed
    ]

    output_dir = r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp'
    output_filename = 'docxtp4.docx'

    merge_docx(input_files, output_dir, output_filename)
