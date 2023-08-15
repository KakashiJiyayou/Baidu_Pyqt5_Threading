from docx import Document
from docx.oxml import OxmlElement
import os
from docx.shared import Inches

def extract_and_replace_images(doc, images, merged_doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for rel, image_path in images.items():
                placeholder = f"IMAGE_SOURCEWILLBEHERE:{rel}"
                print ( "we are here 1st position" ,run.text)
                if placeholder in run.text:
                    print("we are here 2nd position ")
                    run.text = run.text.replace(placeholder, "")
                    run_element = run._r
                    drawing = OxmlElement('w:drawing')
                    inline = OxmlElement('wp:inline')
                    extent = OxmlElement('wp:extent')
                    extent.set('cx', '4000000')
                    extent.set('cy', '3000000')
                    effect_extent = OxmlElement('wp:effectExtent')
                    effect_extent.set('l', '0')
                    effect_extent.set('t', '0')
                    effect_extent.set('r', '0')
                    effect_extent.set('b', '0')
                    doc_pr = OxmlElement('wp:docPr')
                    doc_pr.set('id', '1')
                    doc_pr.set('name', 'Image 1')
                    c_nv_pr = OxmlElement('wp:cNvPr')
                    c_nv_pr.set('id', '2')
                    c_nv_pr.set('name', 'Picture 1')
                    c_nv_pr.set('descr', image_path)
                    blip_fill = OxmlElement('a:blip')
                    blip_fill.set('r:embed', rel)
                    stretch = OxmlElement('a:stretch')
                    fill_rect = OxmlElement('a:fillRect')
                    stretch.append(fill_rect)
                    blip_fill.append(stretch)
                    blip = OxmlElement('a:blipFill')
                    blip.append(blip_fill)
                    src_rect = OxmlElement('a:srcRect')
                    src_rect.set('t', '0')
                    src_rect.set('r', '0')
                    src_rect.set('b', '0')
                    src_rect.set('l', '0')
                    stretch_info = OxmlElement('a:stretch')
                    stretch_info.append(src_rect)
                    blip.append(stretch_info)
                    inline.append(extent)
                    inline.append(effect_extent)
                    inline.append(doc_pr)
                    inline.append(c_nv_pr)
                    inline.append(blip)
                    drawing.append(inline)
                    run_element.append(drawing)

                    # Add the image to merged_doc
                    image_ext = os.path.splitext(image_path)[1][1:]
                    merged_doc.add_picture(image_path, width=Inches(2), height=Inches(1.5))


def merge_docx_files(doc_paths, output_path, index):
    merged_doc = Document()

    # Create a directory for saving the extracted images
    image_output_folder = os.path.join(output_path, 'images')
    os.makedirs(image_output_folder, exist_ok=True)

    for doc_path in doc_paths:
        doc = Document(doc_path)
        images = {}  # Dictionary to store image placeholders and their paths

        # Modify the document to add placeholders for images
        for paragraph in doc.paragraphs:
            for rel in doc.part.rels:
                if 'image' in doc.part.rels[rel].target_ref:
                    placeholder = f"IMAGE_SOURCEWILLBEHERE:{rel}"
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, placeholder)

        # Extract images from the source document and store their paths
        for rel in doc.part.rels:
            if 'image' in doc.part.rels[rel].target_ref:
                image_part = doc.part.rels[rel].target_part
                image_ext = image_part.content_type.split('/')[-1]
                image_name = f"image_{rel}.{image_ext}"
                image_path = os.path.join(image_output_folder, image_name)

                with open(image_path, 'wb') as image_file:
                    image_file.write(image_part.blob)

                images[rel] = image_path

            # Replace image placeholders and add images to merged_doc
        extract_and_replace_images(doc, images, merged_doc)

        for paragraph in doc.paragraphs:
            merged_doc.add_paragraph(paragraph.text)

    # Save the merged document with replaced images
    # Step 4: Save the new DOCX document with replaced image placeholders
    new_docx_name = f"mxt{index}.docx"  # Change the naming convention as needed
    new_docx_path = os.path.join(output_path, new_docx_name)
    merged_doc.save(new_docx_path)
    merged_doc.save(new_docx_path)
    print(f"Merged documents with images saved to {new_docx_path}")


# List of input processed DOCX files
processed_docx_paths = [
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp1.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp2.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp3.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp4.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx1.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx2.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx3.docx',
    r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx4.docx',
    # Add more paths as needed
]

# Output path for the merged document
merged_output_path = r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp'


index = 10
# Call the function to merge the processed documents
merge_docx_files(processed_docx_paths, merged_output_path, index)
