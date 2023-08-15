from docx import Document
from pptx import Presentation
import os
import shutil
def extract_images(docx_path, output_dir):
    doc = Document(docx_path)
    image_paths = []

    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            img_rel = doc.part.rels[rel]
            img_path = os.path.join(output_dir, os.path.basename(img_rel.target_part.target_ref))
            img_data = doc.part.related_parts[img_rel.target_part]
            with open(img_path, "wb") as img_file:
                img_file.write(img_data.blob)
            image_paths.append(img_path)

    return image_paths

def merge_docx(input_files, output_dir, output_filename):
    # Create a new document to hold the merged content
    merged_doc = Document()

    for input_file in input_files:
        # Open the source document
        source_doc = Document(input_file)

        for element in source_doc.element.body:
            # Append paragraphs, tables, and images to the merged document
            merged_doc.element.body.append(element)

    # Save the merged document
    merged_path = os.path.join(output_dir, output_filename)
    merged_doc.save(merged_path)

if __name__ == "__main__":
    input_files = [

        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp4.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx4.docx',
        # Add more paths as needed
    ]

    output_dir = r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp'
    output_filename = 'mo1.docx'

    merge_docx(input_files, output_dir, output_filename)
