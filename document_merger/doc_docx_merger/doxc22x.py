from docxcompose.composer import Composer
from docx import Document

import os
def merge_documents(doc_paths, output_path):
    master_doc = Document()
    master_doc.add_paragraph()
    composer = Composer(master_doc)

    for doc_path in doc_paths:
        try:
            doc = Document(doc_path)
            composer.append(doc)
        except:
            pass

    composer.save(output_path)




# List of paths to the documents you want to merge
input_files = [
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\安全员于娜娜0725.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\附件2：2023 宁波市国际学生汉语口语大赛报名表.docx',
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\CCTV.docx",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\doc2.docx",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\doc3.docx"
    ]

# Output file name for the merged document
output_file = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp\tue15.docx"

merge_documents(input_files, output_file)