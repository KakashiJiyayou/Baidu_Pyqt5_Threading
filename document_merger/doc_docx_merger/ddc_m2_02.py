import os
from docx import Document

def apply_post_merge_formatting(doc):
    for paragraph in doc.paragraphs:
        # Remove any paragraph numbering
        paragraph.clear_numbering()

        # Apply consistent paragraph style if desired
        # paragraph.style = 'Normal'

def merge_docx(input_files, output_dir, output_filename):
    # Load the content of the first document as the starting point
    merged_doc = Document(input_files[0])

    # Remove the first element (body) as it's already added to merged_doc
    # del merged_doc.element.body[0]

    for input_file in input_files[1:]:
        docx_content = Document(input_file)
        for element in docx_content.element.body:
            new_paragraph = merged_doc.add_paragraph()
            for run in element.runs:
                new_run = new_paragraph.add_run(run.text)
                new_run.bold = run.bold
                # Add more run propertie

    # apply_post_merge_formatting(merged_doc)
    # Save the merged document
    merged_path = os.path.join(output_dir, output_filename)
    merged_doc.save(merged_path)


if __name__ == "__main__":
    input_files = [
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\temp4.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx1.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx2.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx3.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\tempx4.docx',
        # Add more paths as needed
    ]

    output_dir = r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp'
    output_filename = 'docxtp4.docx'

    merge_docx(input_files, output_dir, output_filename)
