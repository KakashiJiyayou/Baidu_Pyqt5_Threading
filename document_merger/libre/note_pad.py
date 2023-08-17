from docxcompose.composer import Composer
from docx import Document
import os
import subprocess

def convert_to_docx(input_doc, output_dir):
    try:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        output_docx = os.path.join(output_dir, os.path.splitext(os.path.basename(input_doc))[0] + ".docx")
        soffice_path = r"E:\Installed\LibreOffice\program\soffice.exe"

        subprocess.run([soffice_path, "--headless", "--convert-to", "docx", input_doc, "--outdir", output_dir],
                       check=True)

        print(f"Conversion successful: {input_doc} -> {output_docx}")

        return output_docx
    except subprocess.CalledProcessError as e:
        print("Conversion failed:", e)
        return None

def merge_documents(doc_paths, output_path):
    master_doc = Document()
    master_doc.add_paragraph()
    composer = Composer(master_doc)

    for doc_path in doc_paths:
        try:
            doc = Document(doc_path)
            composer.append(doc)
        except:
            pass

    composer.save(output_path)

def main():
    input_files = [
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\CCTV.doc",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\doc2.docx",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\doc3.docx",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\doc4.doc"
    ]

    # Convert all input documents to .docx format
    converted_documents = []
    for input_doc in input_files:
        if input_doc.lower().endswith('.doc'):
            converted_doc = convert_to_docx(input_doc, output_dir="path/to/temp_folder")
            if converted_doc:
                converted_documents.append(converted_doc)
        elif input_doc.lower().endswith('.docx'):
            converted_documents.append(input_doc)

    # Output file name for the merged document
    output_merged_doc = "path/to/output/merged_document.docx"

    # Merge .docx files
    merge_documents(converted_documents, output_merged_doc)

if __name__ == "__main__":
    main()
