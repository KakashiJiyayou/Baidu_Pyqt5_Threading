from docxcompose.composer import Composer
from docx import Document
import os
import subprocess


def convert_doc_to_docx(input_docs, output_dir):
    new_docx = []
    try:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        for input_doc in input_docs:
            output_docx = os.path.join(output_dir, os.path.splitext(os.path.basename(input_doc))[0] + ".docx")
            soffice_path = r"E:\Installed\LibreOffice\program\soffice.exe"

            subprocess.run([soffice_path, "--headless", "--convert-to", "docx", input_doc, "--outdir", output_dir],
                           check=True)

            print(f"Conversion successful: {input_doc} -> {output_docx}")

            new_docx.append(output_docx)
    except subprocess.CalledProcessError as e:
        print("Conversion failed:", e)

    return new_docx


def merge_documents(doc_paths, output_path):
    master_doc = Document()
    master_doc.add_paragraph()
    composer = Composer(master_doc)

    for doc_path in doc_paths:
        try:
            doc = Document(doc_path)
            composer.append(doc)
        except:
            pass

    composer.save(output_path)


def main():
    input_docs_list = [
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\CCTV.doc",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\doc2.doc",
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\doc3.doc"
    ]
    output_docx_dir = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\temp"

    # Convert .doc to .docx
    converted_docx_list = convert_doc_to_docx(input_docs_list, output_docx_dir)

    # Output file name for the merged document
    output_merged_doc = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\temp\tue16_1.docx"

    # Merge .docx files
    merge_documents(converted_docx_list, output_merged_doc)


if __name__ == "__main__":
    main()
