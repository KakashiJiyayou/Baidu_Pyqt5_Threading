from docxcompose.composer import Composer
from docx import Document

import os
import  time


def merge_documents(doc_paths, output_path):
    doc = Document(doc_paths[0])
    composer  = Composer (doc )

    temp_folder = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\temp"

    for doc_path in doc_paths[1:]:
        try:
            # Open and read the original document
            docc = Document(doc_path)

            # # Save a temporary copy in the temp folder
            # temp_doc_path = os.path.join(temp_folder, os.path.basename(doc_path))
            # doc.save(temp_doc_path)
            #
            # # Close the document
            # del doc  # This will close the document
            #
            # time.sleep(4)
            #
            # # Reopen the temporary copy
            # temp_doc = Document(temp_doc_path)
            composer.append(docc)

            # Uncomment the following line to keep the temporary doc for inspection
            # os.remove(temp_doc_path)

        except:
            pass
    composer.save(output_path)




# List of paths to the documents you want to merge
input_files = [
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\ok_new.docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\big.docx',
        r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\docx\ok.docx',
        # r'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\doc_docx_merger\docx\附件2：2023 宁波市国际学生汉语口语大赛报名表.docx',
        r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\temp\CCTV.docx",
        # r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\temp\doc2.docx",
        # r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\temp\doc3.docx"
    ]

for item in input_files:
    print ("item ", item)
# Output file name for the merged document
output_file = r"E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\libre\libre10.docx"

merge_documents(input_files, output_file)