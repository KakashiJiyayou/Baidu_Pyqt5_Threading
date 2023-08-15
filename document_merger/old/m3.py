from docx import Document
from PIL import Image
import os
import tempfile

import win32com.client
from docxcompose.composer import Composer
import traceback


def add_files_to_doc(   file_list, folder_path, output_filename, output_folder):
    doc = Document()
    composer = Composer(doc)

    print("File list ", file_list)

    for filename in file_list:
        file_path = os.path.join(folder_path, filename)
        print(" add_files_to_doc file path ", file_path)

        try:
            if filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                # Process image content and add to doc
                pass

            elif filename.lower().endswith('.docx'):
                # Process DOCX content and add to doc
                docx_content = Document(file_path)
                composer.append(docx_content)

            elif filename.lower().endswith('.doc'):
                # Process DOC content and add to docx using a temporary DOCX
                temp_docx_path = convert_doc_to_docx(file_path) #FIXME - use self
                docx_content = Document(temp_docx_path)
                composer.append(docx_content)
                os.remove(temp_docx_path)

        except Exception as e:
            print(f"Error processing {filename}: {e}")
            traceback.print_exc()

    output_path = os.path.join(output_folder, output_filename)
    composer.save(output_path)
    print(f"Document '{output_filename}' saved successfully in '{output_folder}'.")



def convert_doc_to_docx( doc_path):
    temp_docx_path = tempfile.mktemp(suffix='.docx')

    # Use pywin32 to interact with Word for DOC to DOCX conversion
    try:
        word = win32com.client.Dispatch("Word.Application")
        doc = word.Documents.Open(doc_path)
        doc.SaveAs2(temp_docx_path, FileFormat=16)  # 16 represents DOCX format
        doc.Close()
        word.Quit()
    except Exception as e:
        print(f"Error converting {doc_path} to DOCX: {e}")
        return None

    return temp_docx_path

file_list = ['Asif Coder.doc', 'ok0.docx' , "CCTV.doc"]
folder_path = 'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger'
output_filename = 'output.docx'
output_folder = 'E:\Project\Job\GQ\Python\Baidu_downlaod_upload\Baidu_Pyqt5_Threading\document_merger\est'

add_files_to_doc(file_list, folder_path, output_filename, output_folder)